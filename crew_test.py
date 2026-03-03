# -*- coding: utf-8 -*-
"""
需求文档 → 分析问题 → 整理测试点 → 生成测试用例 → 评审优化
支持从文本文件读入需求，跑完整个 Crew 流程并保存结果。
测试用例支持导出：Excel（直接下载）、Google 表格。

用法:
  export GEMINI_API_KEY=你的key
  python crew_test.py -f demand.txt              # 指定需求文件
  python crew_test.py -f demand.txt --no-excel   # 不导出 Excel（默认开启）
  python crew_test.py -f demand.txt --export-sheets # 导出到 Google 表格（需配置 GOOGLE_SHEETS_CREDENTIALS_JSON）
  python crew_test.py -f demand.txt --local      # 本地模式：不调用 Gemini，用占位 LLM 跑完四 Agent 并导出 Excel
"""
from __future__ import annotations

import argparse
import html as html_lib
import json
import os
import re
import sys
from datetime import datetime
from typing import Any

from crewai import Agent, Task, Crew
from langchain_google_genai import ChatGoogleGenerativeAI

try:
    import yaml
except ImportError:
    yaml = None

# ---------------------------------------------------------------------------
# 配置
# ---------------------------------------------------------------------------
DEFAULT_DEMAND = "AB test for live streaming resolution."
DEFAULT_REQUIREMENT_FILE = "demand.txt"
OUTPUT_DIR = "output"
CONFIG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config")
AGENTS_CONFIG_PATH = os.path.join(CONFIG_DIR, "agents.yaml")
PROJECT_MEMORY_PATH = os.path.join(CONFIG_DIR, "project_memory.md")
FAMBASE_MODULES_PATH = os.path.join(CONFIG_DIR, "fambase_modules.yaml")
DOC_FILTER_PATH = os.path.join(CONFIG_DIR, "doc_filter.yaml")

# 已弃用模型映射（gemini-1.5 / 2.0 系列逐步下线，统一收敛到 2.5）
_DEPRECATED_GEMINI_MODEL_MAP = {
    "gemini-1.5-pro": "gemini-2.5-pro",
    "gemini-1.5-flash": "gemini-2.5-flash-lite",
    "gemini-2.0-flash": "gemini-2.5-flash-lite",
}


def _resolve_gemini_model(model: str) -> str:
    """将已弃用模型名映射到当前可用模型。"""
    m = (model or "").strip().lower()
    return _DEPRECATED_GEMINI_MODEL_MAP.get(m, m or "gemini-2.5-flash-lite")


def _build_gemini_llm(
    model_name: str,
    gemini_api_key: str,
    cached_content_name: str = "",
) -> ChatGoogleGenerativeAI:
    """构建 Gemini LLM；若 SDK 支持则挂载 cached_content。"""
    base_kwargs = {
        "model": model_name,
        "google_api_key": gemini_api_key,
        "temperature": 0.4,
    }
    if cached_content_name:
        try:
            return ChatGoogleGenerativeAI(cached_content=cached_content_name, **base_kwargs)
        except TypeError:
            # 兼容不支持 cached_content 参数的版本
            pass
    return ChatGoogleGenerativeAI(**base_kwargs)


def _load_doc_filter_config() -> dict[str, Any]:
    """加载 doc_filter.yaml，用于拉取时排除非产品需求文档。"""
    if yaml is None or not os.path.isfile(DOC_FILTER_PATH):
        return {
            "exclude_title_patterns": [
                r"测试用例|^用例$|^TC-|TC\d",
                r"UI走查|UI 走查|走查清单|走查报告",
                r"进度汇总|进度周报|周报|日报|月报",
                r"会议纪要|会议记录",
                r"bug|缺陷|issue|复盘|评审记录|测试报告",
            ],
            "exclude_content_signals": [
                r"用例ID|用例 id",
                r"预期结果.*操作步骤|操作步骤.*预期结果",
                r"走查项|检查项",
            ],
        }
    try:
        with open(DOC_FILTER_PATH, "r", encoding="utf-8") as f:
            cfg = yaml.safe_load(f) or {}
        return {
            "exclude_title_patterns": cfg.get("exclude_title_patterns") or [],
            "exclude_content_signals": cfg.get("exclude_content_signals") or [],
        }
    except Exception:
        return {"exclude_title_patterns": [], "exclude_content_signals": []}


def is_product_requirement_doc(title: str, content: str) -> tuple[bool, str]:
    """判断是否为产品需求文档。返回 (是否保留, 排除原因)。保留则排除原因为空。
    供单文档导入等处调用。"""
    cfg = _load_doc_filter_config()
    title = (title or "").strip()
    preview = (content or "")[:1500]

    for pat in cfg.get("exclude_title_patterns") or []:
        try:
            if re.search(pat, title, re.I):
                return False, f"标题匹配排除模式: {pat[:30]}…"
        except re.error:
            pass

    for pat in cfg.get("exclude_content_signals") or []:
        try:
            if re.search(pat, preview):
                return False, f"内容匹配排除信号: {pat[:30]}…"
        except re.error:
            pass

    return True, ""


def load_agents_config(path: str | None = None) -> dict[str, Any]:
    """从 YAML 加载 Agent 与 Task 定义。
    - 若 path 为空则使用 AGENTS_CONFIG_PATH；
    - 若文件不存在、PyYAML 未安装或解析失败则返回空 dict（由上层做友好提示）。"""
    if yaml is None:
        return {}
    p = path or AGENTS_CONFIG_PATH
    if not os.path.isfile(p):
        return {}
    try:
        with open(p, "r", encoding="utf-8") as f:
            return yaml.safe_load(f) or {}
    except Exception:
        # 解析失败时让调用方展示“未找到或解析失败，可在此编辑并保存”等友好提示
        return {}


def _build_crew_from_config(
    config: dict[str, Any],
    llm: Any,
    project_context: str = "",
    stream: bool = False,
) -> Crew:
    """根据 config（来自 load_agents_config）构建 Crew。project_context 会替换 task description 中的 {project_context}。"""
    agents_cfg = config.get("agents") or []
    tasks_cfg = config.get("tasks") or []
    agent_map: dict[str, Agent] = {}
    for a in agents_cfg:
        aid = a.get("id") or ""
        agent_map[aid] = Agent(
            role=a.get("role", ""),
            goal=a.get("goal", ""),
            backstory=(a.get("backstory") or "").strip(),
            llm=llm,
            verbose=True,
        )
    task_map: dict[str, Task] = {}
    for t in tasks_cfg:
        tid = t.get("id") or ""
        agent_id = t.get("agent_id") or ""
        agent = agent_map.get(agent_id)
        if not agent:
            continue
        desc = (t.get("description") or "").strip()
        desc = desc.replace("{project_context}", project_context.strip())
        ctx_ids = t.get("context") or []
        context_tasks = [task_map[cid] for cid in ctx_ids if cid in task_map]
        task_map[tid] = Task(
            description=desc,
            expected_output=(t.get("expected_output") or "").strip(),
            agent=agent,
            context=context_tasks if context_tasks else None,
        )
    agents_ordered = []
    for a in agents_cfg:
        aid = a.get("id")
        if aid and agent_map.get(aid):
            agents_ordered.append(agent_map[aid])
    tasks_ordered = [task_map[t["id"]] for t in tasks_cfg if t.get("id") in task_map]
    return Crew(
        agents=agents_ordered,
        tasks=tasks_ordered,
        verbose=True,
        stream=stream,
    )


def _run_crew_sequential(
    config: dict[str, Any],
    project_context: str,
    llm_demand: str,
    stream: bool = False,
) -> tuple[str, list[dict[str, str]]]:
    """按 task1→task2→task3→task4 顺序执行，每次将上游输出通过 inputs 注入下游。返回 (最终输出, step_outputs)。"""
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        raise ValueError("ERROR: GEMINI_API_KEY 未设置。")
    model_name = _resolve_gemini_model(os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite"))
    cached_content_name = ""
    try:
        from context_cache_service import refresh_context_cache_if_needed

        cached_content_name = refresh_context_cache_if_needed(
            project_context=project_context or "",
            gemini_key=gemini_api_key,
            gemini_model=model_name,
        )
    except ImportError:
        cached_content_name = ""
    llm = _build_gemini_llm(model_name, gemini_api_key, cached_content_name)
    tasks_cfg = config.get("tasks") or []
    if not tasks_cfg:
        return "", []

    agents_cfg = config.get("agents") or []
    agent_map: dict[str, Agent] = {}
    for a in agents_cfg:
        aid = a.get("id") or ""
        agent_map[aid] = Agent(
            role=a.get("role", ""),
            goal=a.get("goal", ""),
            backstory=(a.get("backstory") or "").strip(),
            llm=llm,
            verbose=True,
        )

    proj_ctx = (project_context or "").strip()
    outputs: dict[str, str] = {}
    step_outputs: list[dict[str, str]] = []

    for t in tasks_cfg:
        tid = t.get("id") or ""
        agent_id = t.get("agent_id") or ""
        agent = agent_map.get(agent_id)
        if not agent:
            continue

        desc = (t.get("description") or "").strip()
        desc = desc.replace("{project_context}", proj_ctx)

        task_obj = Task(
            description=desc,
            expected_output=(t.get("expected_output") or "").strip(),
            agent=agent,
            context=None,
        )
        # 顺序执行模式不需要每个任务单独流式：流式分支会导致 out_str 变成
        # "<generator object ...>" 或对象表示字符串，使 task3 输出丢失、Excel 无法导出。
        # 统一使用非流式执行，通过 .raw 可靠获取每步完整输出。
        crew = Crew(agents=[agent], tasks=[task_obj], verbose=True)

        inputs: dict[str, str] = {"prd_content": llm_demand}
        if outputs.get("task1"):
            inputs["task1_output"] = outputs["task1"]
        if outputs.get("task2"):
            inputs["task2_output"] = outputs["task2"]
        if outputs.get("task3"):
            inputs["task3_output"] = outputs["task3"]

        result = crew.kickoff(inputs=inputs)
        out_str = str(getattr(result, "raw", result))

        out_str = (out_str or "").strip()
        outputs[tid] = out_str
        step_outputs.append({
            "task": tid,
            "agent": agent_id,
            "content": out_str,
        })

    last_tid = tasks_cfg[-1].get("id") if tasks_cfg else ""
    return outputs.get(last_tid, ""), step_outputs


# ---------------------------------------------------------------------------
# 需求加载（不依赖 API Key，可单独自检）
# ---------------------------------------------------------------------------


def desensitize_for_llm(text: str) -> str:
    """对传给大模型的文本做基础脱敏处理，掩盖明显的账号、邮箱等敏感信息。
    规则可后续通过配置扩展，这里实现一版保守的默认策略。"""
    if not text:
        return ""
    out = str(text)
    # 掩盖邮箱
    out = re.sub(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "[EMAIL_MASKED]", out)
    # 掩盖连续 16 位及以上数字（如可能的卡号）
    out = re.sub(r"\b\d{16,}\b", "[NUMBER_MASKED]", out)
    return out


def load_demand(file_path: str | None) -> str:
    """从文件或环境变量或默认值加载需求文本。优先：文件 > 环境变量 DEMAND > 默认字符串。"""
    if file_path and os.path.isfile(file_path):
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read().strip()
        if not content:
            print(f"警告: 文件 {file_path} 为空，使用默认需求。", file=sys.stderr)
            return os.getenv("DEMAND", DEFAULT_DEMAND)
        return content
    return os.getenv("DEMAND", DEFAULT_DEMAND)


# ---------------------------------------------------------------------------
# Crew 构建（延迟初始化：仅在跑流程时检查 Key 并创建 LLM/Agents）
# ---------------------------------------------------------------------------


def _get_crew_with_config(
    agents_config: dict[str, Any],
    project_context: str = "",
    stream: bool = False,
) -> Crew:
    """使用 YAML 配置构建 Crew（不缓存），并注入 project_context。"""
    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        raise ValueError("ERROR: GEMINI_API_KEY 未设置。")
    model_name = _resolve_gemini_model(os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite"))
    cached_content_name = ""
    try:
        from context_cache_service import refresh_context_cache_if_needed

        cached_content_name = refresh_context_cache_if_needed(
            project_context=project_context or "",
            gemini_key=gemini_api_key,
            gemini_model=model_name,
        )
    except ImportError:
        cached_content_name = ""
    llm = _build_gemini_llm(model_name, gemini_api_key, cached_content_name)
    return _build_crew_from_config(agents_config, llm, project_context, stream=stream)


def chat_with_document_agent(
    user_message: str,
    document_context: str,
    project_context: str = "",
    agents_config_path: str | None = None,
) -> str:
    """与产品文档管理 Agent（文档分析师）沟通：基于文档内容回答用户问题。
    用于验证 Agent 对文档的理解，或对文档进行问答。
    - document_context: 需求文档内容（来自项目记忆或手动粘贴）
    - project_context: 项目记忆摘要（可选）
    """
    if not user_message or not user_message.strip():
        return "请输入你的问题。"
    if not document_context or not document_context.strip():
        return "请先提供文档内容：在「与文档 Agent 沟通」页选择「上次运行的需求」或「项目记忆」，或手动粘贴文档。"

    gemini_api_key = os.getenv("GEMINI_API_KEY")
    if not gemini_api_key:
        raise ValueError("GEMINI_API_KEY 未设置，请先配置。")

    model_name = _resolve_gemini_model(os.getenv("GEMINI_MODEL", "gemini-2.5-flash-lite"))
    llm = _build_gemini_llm(model_name, gemini_api_key)

    config = load_agents_config(agents_config_path or AGENTS_CONFIG_PATH)
    agents_cfg = config.get("agents") or []
    doubt_cfg = next((a for a in agents_cfg if a.get("id") == "doubt_agent"), None)
    if not doubt_cfg:
        doubt_cfg = agents_cfg[0] if agents_cfg else None

    if doubt_cfg:
        doc_agent = Agent(
            role=doubt_cfg.get("role", "Document Analyst"),
            goal=doubt_cfg.get("goal", "理解并分析产品需求文档"),
            backstory=(doubt_cfg.get("backstory") or "").strip() + "\n\n你现在需要基于文档内容，直接回答用户的问题。回答要准确、简洁、基于文档事实。",
            llm=llm,
            verbose=True,
        )
    else:
        doc_agent = Agent(
            role="Document Analyst",
            goal="理解产品需求文档并准确回答用户问题",
            backstory="你是资深需求分析师，熟悉产品文档。根据文档内容回答用户问题，回答要基于文档事实。",
            llm=llm,
            verbose=True,
        )

    proj_ctx = f"\n\n【项目背景】\n{project_context}" if project_context and project_context.strip() else ""
    # 入参脱敏：对文档内容做掩码后再传给 Agent，避免直接暴露敏感信息
    safe_doc = desensitize_for_llm(document_context)
    doc_preview = safe_doc[:60000] + ("..." if len(safe_doc) > 60000 else "")

    task = Task(
        description="""你是一位产品文档管理专家。用户会向你提问，你需要基于下方提供的【需求文档】内容回答。

要求：
- 回答必须基于文档事实，不确定时明确说明
- 若文档中无相关信息，如实告知
- 回答简洁清晰，可直接用于与产品/研发沟通

【需求文档】
{doc}

【用户问题】
{question}
{project}
""".replace("{doc}", doc_preview).replace("{question}", user_message.strip()).replace("{project}", proj_ctx),
        expected_output="基于文档的准确回答",
        agent=doc_agent,
    )
    crew = Crew(agents=[doc_agent], tasks=[task], verbose=True)
    result = crew.kickoff()
    out = getattr(result, "raw", result)
    return str(out).strip() if out else ""


def parse_test_cases_file(file_or_path) -> tuple[str, int]:
    """从 Excel/CSV/TXT 解析测试用例为可读文本，供导入到项目记忆。返回 (content, rows_count)。"""
    if hasattr(file_or_path, "read"):
        data = file_or_path.read()
        name = getattr(file_or_path, "name", "") or ""
        if isinstance(data, str):
            data = data.encode("utf-8")
    else:
        with open(file_or_path, "rb") as f:
            data = f.read()
        name = os.path.basename(str(file_or_path))
    ext = (name or "").lower().split(".")[-1] if "." in (name or "") else ""
    rows_count = 0
    lines: list[str] = []
    if ext in ("xlsx", "xls"):
        try:
            import openpyxl
            from io import BytesIO
            wb = openpyxl.load_workbook(BytesIO(data), read_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    cells = [str(c or "").strip() for c in (row or [])]
                    if any(cells):
                        lines.append(" | ".join(cells))
                        rows_count += 1
            wb.close()
        except ImportError:
            raise ValueError("解析 Excel 需安装 openpyxl：pip install openpyxl")
    elif ext == "csv" or (not ext and b"," in data[:500]):
        import csv
        from io import StringIO
        text = data.decode("utf-8-sig", errors="replace")
        for row in csv.reader(StringIO(text)):
            if row:
                lines.append(" | ".join(str(c or "").strip() for c in row))
                rows_count += 1
    else:
        text = data.decode("utf-8-sig", errors="replace")
        for line in text.splitlines():
            stripped = line.strip()
            if stripped:
                lines.append(stripped)
                rows_count += 1
    content = "\n".join(lines)
    return content, rows_count


# 文件上传大小限制（PRD 安全规范）
_MAX_SINGLE_FILE_BYTES = 10 * 1024 * 1024  # 10MB
_MAX_TOTAL_UPLOAD_BYTES = 50 * 1024 * 1024  # 50MB


def _extract_text_from_docx(raw: bytes) -> str:
    """从 Word(.docx) 字节流提取纯文本。需 python-docx。"""
    try:
        from docx import Document
        from io import BytesIO
        doc = Document(BytesIO(raw))
        paras = []
        for p in doc.paragraphs:
            if p.text.strip():
                paras.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    paras.append(" | ".join(cells))
        return "\n".join(paras).strip()
    except ImportError:
        return ""
    except Exception:
        return ""


def parse_uploaded_files(uploaded_files: list) -> tuple[str, str, list[dict[str, Any]]]:
    """解析上传的 .md、.docx（需求文档）与 .xlsx（既有用例）文件，供「文件上传生成用例」使用。
    返回 (需求文档合并文本, 既有用例合并文本, 预览信息列表)。
    需至少 1 个需求文档（.md 或 .docx）；.xlsx 可选。"""
    demand_parts: list[str] = []
    xlsx_parts: list[str] = []
    preview_infos: list[dict[str, Any]] = []
    total_bytes = 0

    for f in uploaded_files or []:
        name = (getattr(f, "name", "") or "").strip()
        ext = name.lower().rsplit(".", 1)[-1] if "." in name else ""
        if ext not in ("md", "docx", "xlsx"):
            continue

        try:
            raw = f.read() if hasattr(f, "read") else getattr(f, "getvalue", lambda: b"")()
        except Exception:
            continue
        if not isinstance(raw, bytes):
            raw = (raw or "").encode("utf-8") if raw else b""

        size = len(raw)
        if size > _MAX_SINGLE_FILE_BYTES:
            continue  # 单文件超限则跳过
        total_bytes += size
        if total_bytes > _MAX_TOTAL_UPLOAD_BYTES:
            break

        if ext == "md":
            text = raw.decode("utf-8-sig", errors="replace").strip()
            demand_parts.append(text)
            preview_infos.append({"name": name, "type": "md", "preview": (text[:200] + "…") if len(text) > 200 else text})
        elif ext == "docx":
            text = _extract_text_from_docx(raw)
            demand_parts.append(text)
            preview_infos.append({"name": name, "type": "docx", "preview": (text[:200] + "…") if len(text) > 200 else text})
        else:  # xlsx
            from io import BytesIO
            _wrapper = BytesIO(raw)
            _wrapper.name = name  # type: ignore
            content, rows = parse_test_cases_file(_wrapper)
            xlsx_parts.append(content)
            preview_infos.append({"name": name, "type": "xlsx", "rows": rows})

    demand_merged = "\n\n---\n\n".join(demand_parts) if demand_parts else ""
    existing_cases = "\n\n".join(xlsx_parts) if xlsx_parts else ""
    return demand_merged, existing_cases, preview_infos


def tables_to_text(tables: list[list[list[str]]]) -> str:
    """将解析出的表格转为可读文本（| 分隔），供归档到项目记忆。"""
    lines: list[str] = []
    for tbl in tables:
        for row in tbl:
            lines.append(" | ".join(str(c or "").strip() for c in row))
        if tbl:
            lines.append("")
    return "\n".join(lines).strip()


def _normalize_table_line(line: str) -> str | None:
    """将疑似表格行规范为 |...| 格式，便于 _parse_markdown_tables 解析。
    若行内含 2 个以上 | 且像表格行，则补全首尾 |；否则返回 None。"""
    s = line.strip()
    if not s or s.count("|") < 2:
        return None
    if not s.startswith("|"):
        s = "| " + s
    if not s.endswith("|"):
        s = s + " |"
    return s


def _extract_table_candidates(text: str) -> list[str]:
    """从文本中提取可能包含表格的候选块（含 ``` 代码块内的内容），增强容错。"""
    candidates: list[str] = []
    normalized = text.replace("｜", "|")
    # 1. 全文作为候选
    candidates.append(normalized)
    # 2. 提取 ```...``` 或 ```markdown|md|text...``` 块内容
    for m in re.finditer(r"```(?:markdown|md|text)?\s*\n([\s\S]*?)```", normalized, re.IGNORECASE):
        block = (m.group(1) or "").strip()
        if block and block.count("|") >= 2:
            candidates.append(block)
    return candidates


def _extract_table_rows_relaxed(text: str) -> list[list[str]]:
    """宽松提取：收集所有含 2 个以上 | 的连续行，组成表格。用于主解析失败时的兜底。"""
    lines = text.replace("｜", "|").split("\n")
    rows: list[list[str]] = []
    current_block: list[list[str]] = []
    for raw in lines:
        line = raw.strip()
        norm = _normalize_table_line(line) or line
        if norm and norm.count("|") >= 2:
            if not norm.startswith("|"):
                norm = "| " + norm
            if not norm.endswith("|"):
                norm = norm + " |"
            parts = norm.split("|")
            cells = [c.strip() for c in parts[1:-1] if c is not None]
            if cells and not all(re.match(r"^[\s\-:]+$", c) for c in cells):
                current_block.append(cells)
        else:
            if len(current_block) >= 2:
                rows = current_block
                break
            current_block = []
    if not rows and current_block and len(current_block) >= 2:
        rows = current_block
    return rows


def _normalize_table_column_count(rows: list[list[str]], max_cols: int = 12) -> list[list[str]]:
    """列数归一：按最大列数补齐空单元格，避免列数不一致导致导出异常。"""
    if not rows:
        return rows
    n = max(len(r) for r in rows)
    n = min(n, max_cols)
    return [(r + [""] * (n - len(r)))[:n] for r in rows]


def _parse_markdown_tables(text: str) -> list[list[list[str]]]:
    """从文本中解析所有 Markdown 表格，返回 [表格1行列表, 表格2行列表, ...]，每表为 [row1, row2, ...]，每行为 [cell, ...]。
    支持稍宽松的表格格式（如缺少首尾 |、``` 代码块包裹），以兼容 LLM 输出差异。
    主解析失败时尝试宽松兜底提取。"""
    normalized_text = text.replace("｜", "|")
    candidates = _extract_table_candidates(normalized_text)
    for candidate in candidates:
        tables = _parse_markdown_tables_inner(candidate)
        if tables:
            return [_normalize_table_column_count(t) for t in tables]
    # 兜底：宽松行提取
    relaxed_rows = _extract_table_rows_relaxed(normalized_text)
    if relaxed_rows:
        normalized_rows = _normalize_table_column_count(relaxed_rows)
        return [normalized_rows]
    return []


def _parse_markdown_tables_inner(text: str) -> list[list[list[str]]]:
    """内部解析逻辑：严格按行解析 Markdown 表格。"""
    tables: list[list[list[str]]] = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = _normalize_table_line(raw) or raw.strip()
        if not line.startswith("|") or not line.endswith("|"):
            i += 1
            continue
        rows = []
        while i < len(lines):
            raw = lines[i]
            normalized = _normalize_table_line(raw)
            line = normalized if normalized else raw.strip()
            if not line.startswith("|") or not line.endswith("|"):
                break
            parts = line.split("|")
            cells = [c.strip() for c in parts[1:-1]]
            # 跳过分隔行（|---|、|:---:|---:| 等，单元格仅含 - : 空格）
            if cells and all(re.match(r"^[\s\-:]+$", c) for c in cells):
                i += 1
                continue
            if cells:
                rows.append(cells)
            i += 1
        if rows:
            tables.append(rows)
    return tables


def _sanitize_cell_for_excel(value: Any) -> str:
    """防止 Excel 公式注入：以 = + - @ 开头的内容前置单引号。
    将 <br> 转为换行符，使 Excel 单元格内正确显示换行。
    仅用于导出到 Excel/CSV 等外部文件，对内存中的原始内容不做修改。"""
    s = "" if value is None else str(value)
    s = s.strip()
    # <br> / <br/> / <br /> 转为换行符，Excel 单元格内可正确显示
    s = re.sub(r"<br\s*/?>", "\n", s, flags=re.I)
    if s and s[0] in ("=", "+", "-", "@"):
        return "'" + s
    return s


def _export_to_excel(tables: list[list[list[str]]], excel_path: str) -> bool:
    """将解析出的表格写入 Excel，若有多个表则合并为同一 sheet（按顺序拼接）。返回是否成功。"""
    try:
        import openpyxl
    except ImportError:
        print("提示: 安装 openpyxl 后可导出 Excel，执行 pip install openpyxl", file=sys.stderr)
        return False
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "测试用例"
    row_num = 0
    for tbl in tables:
        for r in tbl:
            row_num += 1
            for col_num, cell in enumerate(r, 1):
                safe_value = _sanitize_cell_for_excel(cell)
                ws.cell(row=row_num, column=col_num, value=safe_value)
        if tbl:
            row_num += 1
    try:
        wb.save(excel_path)
        return True
    except Exception as e:
        print(f"导出 Excel 失败: {e}", file=sys.stderr)
        return False


def export_tables_to_excel_bytes(tables: list[list[list[str]]]) -> bytes | None:
    """将表格导出为 Excel 二进制，供 st.download_button 使用。不落盘，仅返回 bytes。"""
    try:
        from io import BytesIO
        import openpyxl
    except ImportError:
        return None
    if not tables:
        return None
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "测试用例"
    row_num = 0
    for tbl in tables:
        for r in tbl:
            row_num += 1
            for col_num, cell in enumerate(r, 1):
                safe_value = _sanitize_cell_for_excel(cell)
                ws.cell(row=row_num, column=col_num, value=safe_value)
        if tbl:
            row_num += 1
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _export_to_google_sheets(tables: list[list[list[str]]], title: str) -> str | None:
    """将表格写入 Google Sheets 新建表格。需配置 GOOGLE_SHEETS_CREDENTIALS_JSON（服务账号 JSON 路径或内容）。返回表格 URL 或 None。"""
    if not tables:
        return None
    try:
        import gspread
        from google.oauth2.service_account import Credentials
    except ImportError:
        print("提示: 导出到 Google 表格需安装 gspread 与 google-auth，并配置 GOOGLE_SHEETS_CREDENTIALS_JSON", file=sys.stderr)
        return None
    creds_path_or_json = os.getenv("GOOGLE_SHEETS_CREDENTIALS_JSON")
    if not creds_path_or_json:
        print("提示: 请设置环境变量 GOOGLE_SHEETS_CREDENTIALS_JSON（服务账号 JSON 文件路径或 JSON 字符串）", file=sys.stderr)
        return None
    try:
        if creds_path_or_json.strip().startswith("{"):
            import tempfile
            with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as f:
                f.write(creds_path_or_json)
                path = f.name
            creds = Credentials.from_service_account_file(path, scopes=["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"])
            os.unlink(path)
        else:
            creds = Credentials.from_service_account_file(creds_path_or_json.strip(), scopes=["https://www.googleapis.com/auth/spreadsheets", "https://www.googleapis.com/auth/drive"])
        gc = gspread.authorize(creds)
        sh = gc.create(title)
        sheet = sh.sheet1
        row_offset = 0
        for tbl in tables:
            if tbl:
                sheet.update(f"A{1 + row_offset}", tbl)
                row_offset += len(tbl) + 1
        return sh.url
    except Exception as e:
        print(f"导出到 Google 表格失败: {e}", file=sys.stderr)
        return None


def _sanitize_title_for_filename(title: str) -> str:
    """将需求标题转为安全文件名字符，非法字符替换为下划线。"""
    if not title or not str(title).strip():
        return "需求"
    s = str(title).strip()[:80]
    s = re.sub(r'[<>:"/\\|?*]', "_", s)
    s = re.sub(r"\s+", "_", s)
    return s or "需求"


def _save_result(
    demand: str,
    result_str: str,
    output_dir: str,
    demand_title: str | None = None,
) -> tuple[str, str]:
    """将需求和结果写入 output_dir，返回 (txt 路径, 时间戳)。
    文件命名按 PRD：测试用例_{需求标题简写}_{timestamp}.txt"""
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    title_slug = _sanitize_title_for_filename(demand_title or "")
    base_name = f"测试用例_{title_slug}_{timestamp}"
    out_path = os.path.join(output_dir, f"{base_name}.txt")
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("=== 需求 ===\n\n")
        f.write(demand)
        f.write("\n\n=== 输出 ===\n\n")
        f.write(result_str)
    return out_path, timestamp


def load_project_memory(path: str | None = None) -> str:
    """读取项目记忆文件内容，用于注入 Agent 上下文。"""
    p = path or PROJECT_MEMORY_PATH
    if not os.path.isfile(p):
        return ""
    try:
        with open(p, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception:
        return ""


def load_fambase_modules_for_agent(path: str | None = None) -> str:
    """加载 Fambase 模块定义，格式化为供 Agent 使用的文本。用于主模块/子模块与用例编号（主模块缩写-序号）。"""
    p = path or FAMBASE_MODULES_PATH
    if not os.path.isfile(p) or not yaml:
        return ""
    try:
        with open(p, "r", encoding="utf-8") as f:
            data = yaml.safe_load(f)
        modules = data.get("modules") or []
        lines = ["【Fambase 模块定义】主模块与子模块需严格参照下表，用例编号格式：主模块缩写-序号（如 LIV-0001）"]
        for m in modules:
            main = m.get("main", "")
            abbrev = m.get("abbrev", "")
            pri = m.get("priority", "")
            subs = m.get("sub_modules") or []
            sub_str = "；".join(f"{s.get('name','')}({s.get('abbrev','')})" for s in subs if s.get("name"))
            lines.append(f"- {main} [{abbrev}] {pri}：{sub_str}")
        return "\n".join(lines)
    except Exception:
        return ""


def get_project_context_for_agent(include_store: bool = True) -> str:
    """获取供 Agent 使用的项目上下文：知识库存在时用 project_memory + agent_knowledge，否则沿用原逻辑。"""
    md_ctx = load_project_memory()
    mod_ctx = load_fambase_modules_for_agent()
    if mod_ctx:
        md_ctx = (md_ctx + "\n\n" + mod_ctx).strip() if md_ctx else mod_ctx
    if not include_store:
        return md_ctx
    try:
        from agent_knowledge_service import load_agent_knowledge
        kb = load_agent_knowledge()
        if kb:
            return (md_ctx + "\n\n【Agent 知识库】\n\n" + kb).strip() if md_ctx else kb
    except ImportError:
        pass
    try:
        from memory_store import get_recent_for_agent
        store_ctx = get_recent_for_agent(limit=10, demand_only=True, include_test_cases=True)
        if store_ctx:
            md_ctx = (md_ctx + "\n\n【近期需求与产出记录】\n" + store_ctx).strip()
    except ImportError:
        pass
    return md_ctx


def update_project_memory(addition: str, path: str | None = None, max_chars: int = 20000) -> None:
    """在项目记忆文件末尾追加一段摘要，便于 Agent 保持对项目的熟悉。若超过 max_chars 则只保留尾部。"""
    p = path or PROJECT_MEMORY_PATH
    os.makedirs(os.path.dirname(p), exist_ok=True)
    existing = load_project_memory(p)
    sep = "\n\n---\n\n"
    new_content = (existing + sep + addition.strip()).strip()
    if len(new_content) > max_chars:
        new_content = "...\n\n" + new_content[-max_chars:]
    with open(p, "w", encoding="utf-8") as f:
        f.write(new_content)
    try:
        from context_cache_service import mark_context_cache_dirty

        mark_context_cache_dirty("project_memory_updated")
    except ImportError:
        pass


def run_mock_pipeline(demand: str, output_dir: str = OUTPUT_DIR) -> str:
    """本地模拟流程（不调 API）：四步占位输出，用于自检流程与文件写入。"""
    print("[Mock] 需求摘要:", demand[:200] + ("..." if len(demand) > 200 else ""))
    print("[Mock] 模拟执行: 分析 → 测试点 → 测试用例 → 评审")
    step1 = "1. 需求分析（模拟）\n- 文档完整性待确认\n- 分流比例边界未定义\n- 实验周期与稳定性需明确"
    step2 = "2. 测试点（模拟）\n- 分流正确性\n- 配置下发与回退\n- 指标上报与延迟"
    step3 = "3. 测试用例（模拟）\n- TC01: 用户分流到对照组/实验组\n- TC02: 比例配置异常回退对照组"
    step4 = "4. 评审结论（模拟）\n- 覆盖主要功能与异常路径\n- 建议补充性能与稳定性用例"
    result_str = "\n\n".join([step1, step2, step3, step4])
    print("\n" + "=" * 60 + "\n")
    print(result_str)
    out_path, _ = _save_result(demand, result_str, output_dir)
    print(f"\n结果已保存: {out_path}")
    return result_str


# 本地四 Agent 占位输出（与真实 Crew 输出结构一致，含 Markdown 表格，便于导出 Excel）
_LOCAL_AGENT1 = "【需求风险评估报告 - 占位】\n- 模块：直播分辨率；风险类型：功能/兼容；风险概述：AB test 分流与降级策略需明确；优先级建议：P0；问题描述：分辨率档位与兜底逻辑；建议：与产品确认实验周期与回滚条件。"
_LOCAL_AGENT2 = "【测试点骨架 - 占位】\n- 模块：直播分辨率；场景：AB 分流；测试点ID：TP-001；用例概述：用户被正确分流到对照组/实验组；类型：功能；优先级：P0。\n- 模块：直播分辨率；场景：降级；测试点ID：TP-002；用例概述：异常时回退默认分辨率；类型：异常；优先级：P0。"
_LOCAL_AGENT3_TABLE = """| 用例ID | 模块 | 场景 | 用例概述 | 优先级 | 前置条件 | 操作步骤 | 预期结果 |
| TC-001 | 直播分辨率 | AB分流 | 用户被正确分流到对照组/实验组 | P0 | 用户已登录且实验开启 | 进入直播并查看分辨率档位 | 与实验组配置一致，对照组为默认档位 |
| TC-002 | 直播分辨率 | 降级 | 异常时回退默认分辨率 | P0 | 实验组服务异常或超时 | 触发拉流/切换分辨率 | 自动回退到默认分辨率，无黑屏或卡死 |
| TC-003 | 直播分辨率 | 配置下发 | 服务端配置生效 | P0 | 后端更新实验比例 | 新用户进入直播 | 按新比例分流，老用户会话不变 |
| TC-004 | 直播分辨率 | 边界 | 不支持分辨率时兜底 | P1 | 设备仅支持 360p | 进入直播 | 展示 360p 或友好提示，不崩溃 |"""
_LOCAL_AGENT4 = "✅ 审查通过。用例覆盖分流、降级、配置与边界，表头与字段符合 Fambase 交付宪法。（本地占位模式，未调用 Gemini）"


def run_local_crew_pipeline(
    demand: str,
    output_dir: str = OUTPUT_DIR,
    export_excel: bool = True,
    export_sheets: bool = False,
) -> str:
    """不调用 Gemini：按四 Agent 顺序产出占位结果（含 Markdown 表格），并复用保存/导出逻辑。"""
    print("需求摘要:", demand[:200] + ("..." if len(demand) > 200 else ""))
    print("[Local] 四 Agent 占位模式，不调用 Gemini；结果可正常导出 Excel。")
    print()
    result_str = "\n\n".join([_LOCAL_AGENT1, _LOCAL_AGENT2, _LOCAL_AGENT3_TABLE, _LOCAL_AGENT4])
    print("=" * 60)
    print(result_str)
    print("=" * 60)
    out_path, timestamp = _save_result(demand, result_str, output_dir, demand_title=None)
    print(f"\n结果已保存: {out_path}")

    tables = _parse_markdown_tables(result_str)
    if tables:
        if export_excel:
            title_slug = _sanitize_title_for_filename(None)
            excel_path = os.path.join(output_dir, f"测试用例_{title_slug}_{timestamp}.xlsx")
            if _export_to_excel(tables, excel_path):
                print(f"Excel 已导出: {excel_path}")
        if export_sheets:
            sheets_url = _export_to_google_sheets(tables, f"测试用例_{timestamp}")
            if sheets_url:
                print(f"Google 表格已创建: {sheets_url}")
    return result_str


def run_pipeline(
    demand: str,
    output_dir: str = OUTPUT_DIR,
    mock: bool = False,
    local: bool = False,
    export_excel: bool = True,
    export_sheets: bool = False,
    demand_title: str | None = None,
    agents_config_path: str | None = None,
    project_context: str | None = None,
    return_details: bool = False,
) -> str | dict[str, Any]:
    """跑完整个流程，并把结果写入 output_dir。可同时导出 Excel / Google 表格。
    若 return_details=True，返回 dict：result_str, step_outputs, excel_path, sheets_url, timestamp；否则返回 result_str。"""
    if mock:
        return run_mock_pipeline(demand, output_dir)
    if local:
        return run_local_crew_pipeline(
            demand,
            output_dir=output_dir,
            export_excel=export_excel,
            export_sheets=export_sheets,
        )

    use_config = bool(agents_config_path or os.path.isfile(AGENTS_CONFIG_PATH))
    proj_ctx = project_context if project_context is not None else get_project_context_for_agent()
    stream = return_details

    # 入参脱敏：仅对传入大模型的内容做掩码处理，本地保存仍保留原始需求文本
    llm_demand = desensitize_for_llm(demand)
    step_outputs: list[dict[str, str]] = []
    sheets_url: str | None = None
    excel_path: str | None = None

    if use_config and yaml:
        config = load_agents_config(agents_config_path or AGENTS_CONFIG_PATH)
        if not config or not config.get("agents") or not config.get("tasks"):
            raise ValueError("Agent configuration is missing or invalid. Please check your agents.yaml file.")
        # 采用顺序执行 + inputs 占位符注入（见 docs/四Agent任务编排-开发实施文档）
        result_str, step_outputs = _run_crew_sequential(
            config,
            proj_ctx,
            llm_demand,
            stream=stream,
        )
        if not stream:
            print("需求摘要:", demand[:200] + ("..." if len(demand) > 200 else ""))
            print()
            print("\n" + "=" * 60 + "\n")
            print(result_str)
    else:
        raise ValueError("Could not load agent configuration from YAML. Please ensure agents.yaml is present and correct.")

    out_path, timestamp = _save_result(demand, result_str, output_dir, demand_title)
    if not stream:
        print(f"\n结果已保存: {out_path}")

    tables = _parse_markdown_tables(result_str)
    if tables:
        if export_excel:
            title_slug = _sanitize_title_for_filename(demand_title or "")
            excel_path = os.path.join(output_dir, f"测试用例_{title_slug}_{timestamp}.xlsx")
            if _export_to_excel(tables, excel_path):
                if not stream:
                    print(f"Excel 已导出（可直接下载）: {excel_path}")
        if export_sheets:
            sheets_url = _export_to_google_sheets(tables, f"测试用例_{timestamp}")
            if sheets_url and not stream:
                print(f"Google 表格已创建: {sheets_url}")
    else:
        if export_excel or export_sheets:
            if not stream:
                print("提示: 未从输出中解析到 Markdown 表格，未执行表格导出。", file=sys.stderr)
        excel_path = None

    if return_details:
        return {
            "result_str": result_str,
            "step_outputs": step_outputs,
            "excel_path": excel_path,
            "sheets_url": sheets_url,
            "timestamp": timestamp,
            "txt_path": out_path,
        }
    return result_str


def main() -> int:
    parser = argparse.ArgumentParser(
        description="从文字版需求文档跑完分析→测试点→测试用例→评审流程，并保存结果。"
    )
    parser.add_argument(
        "-f",
        "--file",
        default=None,
        metavar="PATH",
        help=f"需求文档路径（文本）。不指定时尝试读当前目录下的 {DEFAULT_REQUIREMENT_FILE}，再回退到环境变量 DEMAND 或内置默认。",
    )
    parser.add_argument(
        "-o",
        "--output-dir",
        default=OUTPUT_DIR,
        metavar="DIR",
        help=f"结果输出目录，默认: {OUTPUT_DIR}",
    )
    parser.add_argument(
        "--mock",
        action="store_true",
        help="本地模拟流程（不调用 Gemini API），用于自检与 CI。",
    )
    parser.add_argument(
        "--local",
        action="store_true",
        help="本地占位模式：不调用 Gemini，按四 Agent 顺序产出占位结果并正常导出 Excel。",
    )
    parser.add_argument(
        "--no-excel",
        action="store_true",
        help="不导出 Excel（默认会从输出中解析表格并导出 .xlsx 到 output/）。",
    )
    parser.add_argument(
        "--export-sheets",
        action="store_true",
        help="导出到 Google 表格（需 GOOGLE_SHEETS_CREDENTIALS_JSON 环境变量）。",
    )
    args = parser.parse_args()

    file_path = args.file
    if file_path is None and os.path.isfile(DEFAULT_REQUIREMENT_FILE):
        file_path = DEFAULT_REQUIREMENT_FILE
    demand = load_demand(file_path)
    if file_path:
        print(f"已从文件加载需求: {os.path.abspath(file_path)}")
    else:
        print("使用环境变量 DEMAND 或内置默认需求")

    try:
        run_pipeline(
            demand,
            output_dir=args.output_dir,
            mock=args.mock,
            local=args.local,
            export_excel=not args.no_excel,
            export_sheets=args.export_sheets,
        )
        return 0
    except Exception as e:
        print(f"执行失败: {e}", file=sys.stderr)
        raise


if __name__ == "__main__":
    sys.exit(main())
